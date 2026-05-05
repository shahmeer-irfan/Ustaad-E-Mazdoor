"""
Convert PROJECT_REPORT.md to PROJECT_REPORT.docx using python-docx.

This is intentionally simple: it walks the markdown, treats lines starting with
"# / ## / ### / ####" as headings and everything else as paragraphs. Tables
(GitHub-flavoured) and fenced code blocks are rendered specially. Mermaid
fences are skipped (Word can't render Mermaid; we mention the file path).
"""
from __future__ import annotations
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(r"C:/Users/shahmeer.irfan/Downloads/Ustaad/Ustaad/PROJECT_REPORT.md")
DST = Path(r"C:/Users/shahmeer.irfan/Downloads/Ustaad/Ustaad/PROJECT_REPORT.docx")

# --- styling helpers ----------------------------------------------------------

def set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BBBBBB")
    pBdr.append(bottom)
    pPr.append(pBdr)


# --- markdown helpers ---------------------------------------------------------

INLINE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
INLINE_BOLD = re.compile(r"\*\*(.+?)\*\*")
INLINE_ITAL = re.compile(r"\*(.+?)\*")
INLINE_CODE = re.compile(r"`([^`]+)`")
HTML_TAG    = re.compile(r"<[^>]+>")


def add_runs(paragraph, text: str) -> None:
    """Render a markdown line into a paragraph with bold/italic/code/link runs."""
    text = HTML_TAG.sub("", text).strip()
    if not text:
        return

    # Resolve links → "label (url)"
    text = INLINE_LINK.sub(r"\1 (\2)", text)

    # Tokenise: split on bold first, then italic, then code, in sequence.
    # Simpler approach: process by repeatedly finding the earliest formatter.
    pos = 0
    pattern = re.compile(
        r"\*\*(?P<bold>.+?)\*\*"
        r"|`(?P<code>[^`]+)`"
        r"|\*(?P<ital>.+?)\*"
    )
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group("bold"):
            r = paragraph.add_run(m.group("bold"));   r.bold   = True
        elif m.group("code"):
            r = paragraph.add_run(m.group("code"));   r.font.name = "Consolas"; r.font.size = Pt(10)
        elif m.group("ital"):
            r = paragraph.add_run(m.group("ital"));   r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Return (rows, next_index). Markdown table starts at lines[start]."""
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        raw = lines[i].strip()
        # skip the header-divider line --- | --- | ---
        if re.match(r"^\|[\s\-:|]+\|$", raw):
            i += 1
            continue
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""  # clear default
            p = cell.paragraphs[0]
            add_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9.5)
                if r_idx == 0:
                    run.bold = True
            if r_idx == 0:
                set_cell_shading(cell, "F2EFEA")
    doc.add_paragraph()


# --- main conversion ----------------------------------------------------------

def build() -> None:
    md = SRC.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = Document()

    # Document defaults
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    in_code = False
    code_buf: list[str] = []
    skip_block = False  # for mermaid / html-only blocks

    i = 0
    while i < len(lines):
        line = lines[i]

        # Code fence handling
        if line.lstrip().startswith("```"):
            if not in_code:
                fence = line.lstrip()[3:].strip().lower()
                in_code = True
                code_buf = []
                skip_block = fence in {"mermaid"}
                i += 1
                continue
            else:
                # close fence
                in_code = False
                if skip_block:
                    p = doc.add_paragraph()
                    r = p.add_run("[Mermaid diagram — view PROJECT_REPORT.md on GitHub for the rendered version]")
                    r.italic = True
                    r.font.size = Pt(10)
                else:
                    p = doc.add_paragraph()
                    for cl in code_buf:
                        run = p.add_run(cl + "\n")
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)
                skip_block = False
                code_buf = []
                i += 1
                continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # HTML <div> blocks — skip the wrapping tags
        if line.strip() in {"<div align=\"center\">", "</div>"}:
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level = min(len(m.group(1)), 4)
            text = HTML_TAG.sub("", m.group(2)).strip()
            text = re.sub(r"^#+\s*", "", text)
            heading = doc.add_heading(level=level)
            run = heading.add_run(text)
            run.font.name = "Calibri"
            if level == 1:
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run.font.size = Pt(28)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            elif level == 2:
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(0xC2, 0x4D, 0x10)  # saffron
            elif level == 3:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            else:
                run.font.size = Pt(12)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^\s*---\s*$", line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables (line starts with |)
        if line.lstrip().startswith("|"):
            rows, next_i = parse_md_table(lines, i)
            add_table(doc, rows)
            i = next_i
            continue

        # Bullet
        bm = re.match(r"^\s*[-*]\s+(.*)", line)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, bm.group(1))
            i += 1
            continue

        # Numbered list
        nm = re.match(r"^\s*\d+\.\s+(.*)", line)
        if nm:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, nm.group(1))
            i += 1
            continue

        # Block quote
        bq = re.match(r"^\s*>\s?(.*)", line)
        if bq:
            p = doc.add_paragraph()
            run = p.add_run(bq.group(1))
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Blank line
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue

        # Default paragraph
        p = doc.add_paragraph()
        add_runs(p, line)
        i += 1

    doc.save(DST)
    msg = f"OK -> {DST}  ({DST.stat().st_size / 1024:.1f} KB)"
    sys.stdout.buffer.write(msg.encode("utf-8") + b"\n")


if __name__ == "__main__":
    try:
        build()
    except Exception as e:
        print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)
